"""
resume_parser.py — Extract plain text from a PDF or DOCX resume and match
skills against the skills taxonomy.

PDF libraries tried in order (first available wins):
  1. pypdf        (pip install pypdf)
  2. PyPDF2       (pip install PyPDF2)
  3. pdfplumber   (pip install pdfplumber)

If none is installed the module will auto-install pypdf via pip.

DOCX requires:
  python-docx    (pip install python-docx)
"""
import re
import logging
import subprocess
import sys

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# PDF library helpers
# ---------------------------------------------------------------------------

def pdf_library_available() -> bool:
    """Return True if at least one PDF parsing library is importable."""
    for lib in ("pypdf", "PyPDF2", "pdfplumber"):
        try:
            __import__(lib)
            return True
        except ImportError:
            continue
    return False


def _auto_install_pypdf() -> None:
    """Install pypdf via pip when no PDF library is present."""
    logger.info("No PDF library found — auto-installing pypdf…")
    try:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "--quiet", "pypdf"],
            timeout=120,
        )
        logger.info("pypdf installed successfully.")
    except subprocess.CalledProcessError as exc:
        raise RuntimeError(
            f"Failed to auto-install pypdf (exit code {exc.returncode}). "
            "Run manually: pip install pypdf"
        ) from exc
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError(
            "pip install pypdf timed out. Run manually: pip install pypdf"
        ) from exc


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def _extract_pdf_text(path: str) -> str:
    """Try available PDF libraries in order and return extracted text."""

    # 1. pypdf (modern, recommended)
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except ImportError:
        pass
    except Exception as exc:
        logger.debug("pypdf extraction failed for %s: %s", path, exc)

    # 2. PyPDF2 (legacy)
    try:
        from PyPDF2 import PdfReader  # type: ignore
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except ImportError:
        pass
    except Exception as exc:
        logger.debug("PyPDF2 extraction failed for %s: %s", path, exc)

    # 3. pdfplumber (highest quality, heavier dependency)
    try:
        import pdfplumber  # type: ignore
        with pdfplumber.open(path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages)
    except ImportError:
        pass
    except Exception as exc:
        logger.debug("pdfplumber extraction failed for %s: %s", path, exc)

    # No library available — auto-install pypdf and retry once
    _auto_install_pypdf()
    try:
        from pypdf import PdfReader  # noqa: F811
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except ImportError as exc:
        raise RuntimeError(
            "pypdf was installed but could not be imported. "
            "Restart the application and try again."
        ) from exc
    except Exception as exc:
        raise RuntimeError(f"pypdf installed but extraction failed: {exc}") from exc


def _extract_docx_text(path: str) -> str:
    """Extract text from a .docx file."""
    try:
        import docx  # type: ignore
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )
    except Exception as exc:
        raise RuntimeError(f"Failed to read DOCX: {exc}")


def extract_resume_text(path: str) -> str:
    """
    Extract plain text from a resume file.
    Supports .pdf and .docx/.doc.
    Raises RuntimeError with an installation hint if a required library is missing.
    """
    if not path:
        raise ValueError("No resume path provided.")

    pl = path.lower()
    if pl.endswith(".pdf"):
        return _extract_pdf_text(path)
    elif pl.endswith((".docx", ".doc")):
        return _extract_docx_text(path)
    else:
        # Unknown extension — try PDF first, then DOCX
        try:
            return _extract_pdf_text(path)
        except Exception:
            return _extract_docx_text(path)


# ---------------------------------------------------------------------------
# Skill matching
# ---------------------------------------------------------------------------

def extract_skills_from_resume(path: str) -> list:
    """
    Parse the resume at *path* and return a list of skills from the taxonomy
    that appear in the resume text.

    Matching rules:
      - Skills ≤ 4 characters: require a word boundary (avoids false positives
        like "Go" matching "going" or "C" matching every sentence).
      - Longer skills: case-insensitive substring match.
    """
    from skills_taxonomy import get_taxonomy

    text = extract_resume_text(path)
    if not text.strip():
        return []

    taxonomy = get_taxonomy()
    found: list = []
    seen: set = set()

    for cat in taxonomy.values():
        for skills in cat.values():
            for skill in skills:
                skill_lower = skill.lower()
                if skill_lower in seen:
                    continue
                try:
                    if len(skill) <= 4:
                        pattern = r'(?<![A-Za-z0-9])' + re.escape(skill) + r'(?![A-Za-z0-9])'
                    else:
                        pattern = re.escape(skill)
                    if re.search(pattern, text, re.IGNORECASE):
                        found.append(skill)
                        seen.add(skill_lower)
                except re.error:
                    if skill_lower in text.lower():
                        found.append(skill)
                        seen.add(skill_lower)

    return found
